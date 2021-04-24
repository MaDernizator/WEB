from docx import Document
from frame.task_generator import TaskGenerator
from frame.task_generator import Pattern
import os
import zipfile
import time
import os
import sys

sys.excepthook = lambda *a: sys.__excepthook__(*a)


class DocGenerator:
    def __init__(self, data, user='anonym'):
        self.name = data['name']
        self.variant_count = int(data['variant_count'])
        self.user = user
        self.patterns = []
        error = True
        for pattern_id in data:
            if pattern_id.isdigit():
                pattern = data[pattern_id]
                if pattern != '0':
                    error = False
                    self.patterns.append(Pattern(pattern_id, int(pattern)))
        if error:
            raise ValueError

    def write_log(self):
        f = open('log/log.txt', 'a')
        data = time.ctime(time.time())
        f.write(f'create {data} | {self.user} | {self.name}\n')
        f.close()

    def directory(self):
        if self.user not in os.listdir('static/generated_documents'):
            os.mkdir(f'static/generated_documents/{self.user}')

    def archive(self):
        dir = zipfile.ZipFile(f'static/generated_documents/{self.user}/{self.name}.zip', 'w')
        for folder, subfolders, files in os.walk(f'static/generated_documents/{self.user}'):
            for file in files:
                if file.endswith('.docx'):
                    dir.write(os.path.join(folder, file),
                              os.path.relpath(os.path.join(folder, file),
                                              f'static/generated_documents/{self.user}'),
                              compress_type=zipfile.ZIP_DEFLATED)

        dir.close()
        for folder, subfolders, files in os.walk(f'static/generated_documents/{self.user}'):
            for file in files:
                if file.endswith('.docx'):
                    os.remove(f'static/generated_documents/{self.user}/{file}')

    def generate_document(self):

        def generate_variant(variant):
            tasks = Document()
            tasks.add_heading(self.name + f' Вариант-{variant}', 0)
            number = 1
            for pattern in self.patterns:
                for _ in range(pattern.get_count()):
                    task_gen = TaskGenerator(pattern.get_pattern())
                    p = tasks.add_paragraph(f'№{number} ' + task_gen.get_text()[0])
                    p.alignment = 3
                    par.add_run(f'\n№{number} - ' + str(task_gen.get_text()[1]))
                    number += 1
            tasks.save(
                f'static/generated_documents/{self.user}/' + self.name + f'_вариант-{variant}.docx')

        self.directory()
        answers = Document()
        for variant in range(1, int(self.variant_count) + 1):
            answers.add_heading(f'Вариант-{variant}', 0)
            par = answers.add_paragraph()
            generate_variant(variant)
        answers.save(f'static/generated_documents/{self.user}/' + self.name + '_ответы.docx')
        self.archive()
        self.write_log()
        return self.name
